from __future__ import annotations as _annotations

import asyncio
import re
import sys
import os
import unicodedata
from contextlib import asynccontextmanager
from dataclasses import dataclass
from typing import List
import json
from docx import Document
from PyPDF2 import PdfReader

import asyncpg
import httpx
# import logfire
import pydantic_core
from openai import AsyncOpenAI
from pydantic import TypeAdapter
from typing_extensions import AsyncGenerator

from pydantic_ai import RunContext
from pydantic_ai.agent import Agent

import logging
from dotenv import load_dotenv

load_dotenv()

# Configure logger to print to stdout
logger = logging.getLogger()
handler = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)
logger.setLevel(logging.INFO)

# # 'if-token-present' means nothing will be sent (and the example will work) if you don't have logfire configured
# logfire.configure(send_to_logfire='if-token-present')
# logfire.instrument_asyncpg()


@dataclass
class Deps:
    openai: AsyncOpenAI
    pool: asyncpg.Pool


agent = Agent('openai:gpt-4o', deps_type=Deps)


@agent.tool
async def retrieve(context: RunContext[Deps], search_query: str) -> str:
    """Retrieve documentation sections based on a search query.

    Args:
        context: The call context.
        search_query: The search query.
    """
    # with logfire.span(
    #     'create embedding for {search_query=}', search_query=search_query
    # ):
    try:
        embedding = await context.deps.openai.embeddings.create(
            input=search_query,
            model='text-embedding-3-small',
        )
    except Exception as e:
        logger.error(f"Embeddings API error for search query '{search_query}': {str(e)}")
        logger.error(f"Request details: model=text-embedding-3-small, input={search_query}")
        raise

    assert (
        len(embedding.data) == 1
    ), f'Expected 1 embedding, got {len(embedding.data)}, doc query: {search_query!r}'
    embedding = embedding.data[0].embedding
    embedding_json = pydantic_core.to_json(embedding).decode()
    rows = await context.deps.pool.fetch(
        'SELECT url, title, content FROM doc_sections ORDER BY embedding <-> $1 LIMIT 8',
        embedding_json,
    )
    return '\n\n'.join(
        f'# {row["title"]}\nDocumentation URL:{row["url"]}\n\n{row["content"]}\n'
        for row in rows
    )


async def run_agent(question: str):
    """Entry point to run the agent and perform RAG based question answering."""
    openai = AsyncOpenAI()
    # logfire.instrument_openai(openai)
    #
    # logfire.info('Asking "{question}"', question=question)

    async with database_connect(False) as pool:
        deps = Deps(openai=openai, pool=pool)
        answer = await agent.run(question, deps=deps)
    # print(answer.data)
    return answer.data


#######################################################
# The rest of this file is dedicated to preparing the #
# search database, and some utilities.                #
#######################################################

# # JSON document from
# # https://gist.github.com/samuelcolvin/4b5bb9bb163b1122ff17e29e48c10992
# DOCS_JSON = (
#     'https://gist.githubusercontent.com/'
#     'samuelcolvin/4b5bb9bb163b1122ff17e29e48c10992/raw/'
#     '80c5925c42f1442c24963aaf5eb1a324d47afe95/logfire_docs.json'
# )
#

def chunk_text(text: str, max_tokens: int = 6000) -> List[str]:
    """Split text into chunks that won't exceed token limit.
    
    Args:
        text (str): Text to split
        max_tokens (int): Maximum tokens per chunk (default 6000 to be safe)
        
    Returns:
        List[str]: List of text chunks
    """
    # Rough estimate: 1 token ≈ 4 characters
    chars_per_chunk = max_tokens * 4
    
    # Split into paragraphs first
    paragraphs = text.split('\n\n')
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_length = len(para)
        
        if current_length + para_length > chars_per_chunk:
            if current_chunk:  # Save current chunk if it exists
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_length = 0
                
        current_chunk.append(para)
        current_length += para_length
        
    if current_chunk:  # Don't forget the last chunk
        chunks.append('\n\n'.join(current_chunk))
        
    return chunks

def build_search_db_from_folder(folder_path: str) -> List[DocsSection]:
    """Build search database entries from documents in a folder.
    
    Args:
        folder_path (str): Path to folder containing documents
        
    Returns:
        List[DocsSection]: List of document sections ready for database insertion
    """
    sections = []
    doc_id = 1
    
    logger.info(f"Scanning folder: {folder_path}")
    files = os.listdir(folder_path)
    logger.info(f"Found {len(files)} files in folder")
    
    for filename in files:
        filepath = os.path.join(folder_path, filename)
        if not os.path.isfile(filepath):
            logger.debug(f"Skipping non-file: {filepath}")
            continue
            
        file_ext = os.path.splitext(filename)[1].lower()
        content = ""
        
        logger.info(f"Processing file: {filepath}")
        
        if file_ext == '.docx':
            logger.debug(f"Parsing DOCX file: {filename}")
            doc = Document(filepath)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            logger.debug(f"Extracted {len(doc.paragraphs)} paragraphs from DOCX")
        elif file_ext == '.pdf':
            logger.debug(f"Parsing PDF file: {filename}")
            reader = PdfReader(filepath)
            content = '\n'.join([page.extract_text() for page in reader.pages])
            logger.debug(f"Extracted {len(reader.pages)} pages from PDF")
        elif file_ext == '.txt':
            logger.debug(f"Reading text file: {filename}")
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            logger.debug(f"Skipping unsupported file type: {filename}")
            continue
            
        # Split content into chunks if it's too large
        content_chunks = chunk_text(content)
        
        # Create a section for each chunk
        for i, chunk in enumerate(content_chunks):
            chunk_title = os.path.basename(filepath)
            if len(content_chunks) > 1:
                chunk_title += f" (Part {i+1}/{len(content_chunks)})"
                
            section = DocsSection(
                id=doc_id,
                parent=None,
                path=os.path.relpath(filepath, start=folder_path),
                level=1,
                title=chunk_title,
                content=chunk
            )
            sections.append(section)
            doc_id += 1
        
    return sections

async def build_search_db(path: str):
    logger.info(f"Starting database build from path: {path}")
    """Build the search database from a URL or local file path.
    
    Args:
        path (str): URL or file path. Can be:
            - HTTP/HTTPS URL
            - file:// URL
            - Absolute or relative filesystem path
    """
    if path.startswith(('http://', 'https://')):
        # Handle remote URLs
        async with httpx.AsyncClient(verify=False) as client:
            response = await client.get(path)
            response.raise_for_status()
        sections = sessions_ta.validate_json(response.content)
    else:
        # Handle local paths (both file:// and regular paths)
        if path.startswith('file://'):
            file_path = path[7:]  # Remove file:// prefix
        else:
            file_path = path
            
        # Convert relative path to absolute path
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
            
        if os.path.isdir(file_path):
            logger.info(f"Processing directory: {file_path}")
            sections = build_search_db_from_folder(file_path)
            logger.info(f"Found {len(sections)} document sections in directory")
        elif file_path.endswith('.json'):
            logger.info(f"Processing JSON file: {file_path}")
            with open(file_path, 'rb') as f:
                content = f.read()
            sections = sessions_ta.validate_json(content)
            logger.info(f"Parsed {len(sections)} sections from JSON")
        else:
            logger.info(f"Processing parent directory of: {file_path}")
            sections = build_search_db_from_folder(os.path.dirname(file_path))

    openai = AsyncOpenAI()
    # logfire.instrument_openai(openai)

    async with database_connect(recreate_db=True) as pool:
        # with logfire.span('create schema'):
        async with pool.acquire() as conn:
            async with conn.transaction():
                await conn.execute(DB_SCHEMA)

        sem = asyncio.Semaphore(10)
        async with asyncio.TaskGroup() as tg:
            for section in sections:
                tg.create_task(insert_doc_section(sem, openai, pool, section))


async def insert_doc_section(
    sem: asyncio.Semaphore,
    openai: AsyncOpenAI,
    pool: asyncpg.Pool,
    section: DocsSection,
) -> None:
    logger.debug(f"Processing section: {section.title} from {section.path}")
    async with sem:
        url = section.url()
        exists = await pool.fetchval('SELECT 1 FROM doc_sections WHERE url = $1', url)
        if exists:
            logger.info('Skipping {url=}', url=url)

        # with logfire.span('create embedding for {url=}', url=url):
        try:
            embedding = await openai.embeddings.create(
                input=section.embedding_content(),
                model='text-embedding-3-small',
            )
        except Exception as e:
            logger.error(f"Embeddings API error for section '{section.title}': {str(e)}")
            logger.error(f"Request details: model=text-embedding-3-small, input length={len(section.embedding_content())}")
            logger.error(f"Section path: {section.path}")
            raise
        assert (
            len(embedding.data) == 1
        ), f'Expected 1 embedding, got {len(embedding.data)}, doc section: {section}'
        embedding = embedding.data[0].embedding
        embedding_json = pydantic_core.to_json(embedding).decode()
        await pool.execute(
            'INSERT INTO doc_sections (url, title, content, embedding) VALUES ($1, $2, $3, $4)',
            url,
            section.title,
            section.content,
            embedding_json,
        )


@dataclass
class DocsSection:
    id: int
    parent: int | None
    path: str
    level: int
    title: str
    content: str

    def url(self) -> str:
        url_path = re.sub(r'\.md$', '', self.path)
        return (
            f'https://logfire.pydantic.dev/docs/{url_path}/#{slugify(self.title, "-")}'
        )

    def embedding_content(self) -> str:
        return '\n\n'.join((f'path: {self.path}', f'title: {self.title}', self.content))


sessions_ta = TypeAdapter(list[DocsSection])


# pyright: reportUnknownMemberType=false
# pyright: reportUnknownVariableType=false
@asynccontextmanager
async def database_connect(
    create_db: bool = False,
    recreate_db: bool = False,
) -> AsyncGenerator[asyncpg.Pool, None]:
    server_dsn, database = (
        'postgresql://postgres:postgres@localhost:54320',
        'pydantic_ai_rag',
    )
    if create_db or recreate_db:
        # with logfire.span('check and create DB'):
        conn = await asyncpg.connect(server_dsn)
        try:
            db_exists = await conn.fetchval(
                'SELECT 1 FROM pg_database WHERE datname = $1', database
            )
            if recreate_db and db_exists:
                # Terminate all connections to the database before dropping
                await conn.execute(
                    f"""
                    SELECT pg_terminate_backend(pg_stat_activity.pid)
                    FROM pg_stat_activity
                    WHERE pg_stat_activity.datname = '{database}'
                    AND pid <> pg_backend_pid()
                    """
                )
                await conn.execute(f'DROP DATABASE {database}')
                db_exists = False
            
            if not db_exists:
                await conn.execute(f'CREATE DATABASE {database}')
        finally:
            await conn.close()

    pool = await asyncpg.create_pool(f'{server_dsn}/{database}')
    try:
        yield pool
    finally:
        await pool.close()


DB_SCHEMA = """
CREATE EXTENSION IF NOT EXISTS vector;

CREATE TABLE IF NOT EXISTS doc_sections (
    id serial PRIMARY KEY,
    url text NOT NULL UNIQUE,
    title text NOT NULL,
    content text NOT NULL,
    -- text-embedding-3-small returns a vector of 1536 floats
    embedding vector(1536) NOT NULL
);
CREATE INDEX IF NOT EXISTS idx_doc_sections_embedding ON doc_sections USING hnsw (embedding vector_l2_ops);
"""


def slugify(value: str, separator: str, unicode: bool = False) -> str:
    """Slugify a string, to make it URL friendly."""
    # Taken unchanged from https://github.com/Python-Markdown/markdown/blob/3.7/markdown/extensions/toc.py#L38
    if not unicode:
        # Replace Extended Latin characters with ASCII, i.e. `žlutý` => `zluty`
        value = unicodedata.normalize('NFKD', value)
        value = value.encode('ascii', 'ignore').decode('ascii')
    value = re.sub(r'[^\w\s-]', '', value).strip().lower()
    return re.sub(rf'[{separator}\s]+', separator, value)


if __name__ == '__main__':
    action = sys.argv[1] if len(sys.argv) > 1 else "build"
    if action == 'build':
        # asyncio.run(build_search_db('file://./data/sample_data.json'))
        asyncio.run(build_search_db("./data/"))
    elif action == 'search':
        if len(sys.argv) == 3:
            q = sys.argv[2]
        else:
            q = 'How do I configure logfire to work with FastAPI?'
        asyncio.run(run_agent(q))
    else:
        print(
            'uv run --extra examples -m pydantic_ai_examples.rag build|search',
            file=sys.stderr,
        )
        sys.exit(1)
